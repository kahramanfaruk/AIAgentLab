from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

from docx import Document as DocxDocument
from pypdf import PdfReader


SUPPORTED_EXTENSIONS = {".pdf", ".txt", ".docx"}


@dataclass(slots=True)
class Document:
    content: str
    source: str
    doc_id: str
    metadata: dict


class DocumentLoader:
    def __init__(self, data_dir: str | Path = "data/raw") -> None:
        self.data_dir = Path(data_dir)

    def load(self) -> list[Document]:
        if not self.data_dir.exists():
            raise FileNotFoundError(f"Data directory not found: {self.data_dir}")

        documents: list[Document] = []
        for path in sorted(self.data_dir.iterdir()):
            if path.is_file() and path.suffix.lower() in SUPPORTED_EXTENSIONS:
                documents.extend(self._load_file(path))
        return documents

    def _load_file(self, path: Path) -> list[Document]:
        suffix = path.suffix.lower()

        if suffix == ".pdf":
            return self._load_pdf(path)
        if suffix == ".txt":
            return [self._load_txt(path)]
        if suffix == ".docx":
            return [self._load_docx(path)]

        return []

    def _load_pdf(self, path: Path) -> list[Document]:
        reader = PdfReader(str(path))
        documents: list[Document] = []

        for page_number, page in enumerate(reader.pages, start=1):
            text = page.extract_text(extraction_mode="layout") or ""
            text = self._normalize_text(text)

            if not text.strip():
                continue

            documents.append(
                Document(
                    content=text,
                    source=str(path),
                    doc_id=f"{path.stem}-page-{page_number}",
                    metadata={
                        "file_name": path.name,
                        "file_type": "pdf",
                        "page_number": page_number,
                    },
                )
            )

        return documents

    def _load_txt(self, path: Path) -> Document:
        text = path.read_text(encoding="utf-8", errors="ignore")
        text = self._normalize_text(text)

        return Document(
            content=text,
            source=str(path),
            doc_id=path.stem,
            metadata={
                "file_name": path.name,
                "file_type": "txt",
            },
        )

    def _load_docx(self, path: Path) -> Document:
        doc = DocxDocument(str(path))
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        text = "\n\n".join(paragraphs)
        text = self._normalize_text(text)

        return Document(
            content=text,
            source=str(path),
            doc_id=path.stem,
            metadata={
                "file_name": path.name,
                "file_type": "docx",
            },
        )

    @staticmethod
    def _normalize_text(text: str) -> str:
        lines = [line.strip() for line in text.splitlines()]
        cleaned = "\n".join(line for line in lines if line)
        return cleaned.strip()


def load_documents(data_dir: str | Path = "data/raw") -> list[Document]:
    return DocumentLoader(data_dir=data_dir).load()


if __name__ == "__main__":
    docs = load_documents()
    print(f"Loaded {len(docs)} documents/pages.")
    for doc in docs[:3]:
        print("-" * 80)
        print(doc.doc_id)
        print(doc.metadata)
        print(doc.content[:500])